#!/usr/bin/env python3
import sys
sys.path.insert(0, '/root/.nvm/versions/node/v24.16.0/lib/python3.13/site-packages')

from docx import Document
import json

doc_path = '/root/.openclaw/media/inbound/翰林光年_对外定位说明书_家长可读版---3ad3135e-a4dc-488c-be38-7da6def746f5.docx'

try:
    doc = Document(doc_path)
    
    print("=" * 70)
    print("翰林光年_对外定位说明书_家长可读版 - 内容提取")
    print("=" * 70)
    
    # Extract business services from document
    print("\n" + "=" * 70)
    print("【业务范围】（从文档中提取）")
    print("=" * 70)
    
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text.strip())
            
    # Extract tables
    print("\n" + "=" * 70)
    print("【表格内容】（如有）")
    print("=" * 70)
    
    for i, table in enumerate(doc.tables):
        print(f"\n表格 {i+1}:")
        for row_idx, row in enumerate(table.rows):
            cell_contents = []
            for cell in row.cells:
                cell_contents.append(cell.text.strip())
            
            if all(cell_contents):  # Only print rows with content
                print(" | ".join(cell_contents))
                
    print("\n" + "=" * 70)
    
except Exception as e:
    print(f"读取文档时出错：{e}")
