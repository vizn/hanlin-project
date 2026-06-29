#!/usr/bin/env python3
"""提取 Word 文档中的文本内容"""
from docx import Document
import sys

# 读取附件的 Word 文件（从 stdin 传递路径）
doc_path = sys.argv[1] if len(sys.argv) > 1 else "/root/.openclaw/workspace/山西代理商版本.docx"

try:
    # 尝试创建文档对象
    print(f"正在读取：{doc_path}")
    
    # 如果文件不存在，从附件路径获取
    import shutil
    if not doc_path.startswith("/"):
        base_name = doc_path.split("/")[-1]
        target_path = f"/root/.openclaw/workspace/{base_name}"
        print(f"将 {doc_path} 复制到工作目录")
        shutil.copy(doc_path, target_path)
        doc_path = target_path
    
    # 读取文档内容
    doc = Document(doc_path)
    
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    tables_text = ""
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            tables_text += " | ".join(cells) + "\n"
    
    # 合并所有内容
    full_text = "\n\n".join(text_parts)
    if tables_text:
        full_text = full_text + f"\n\n{tables_text}"
    
    print(f"\n✅ 成功提取 {len(doc.paragraphs)} 个段落")
    print(f"总字符数：{len(full_text)}")
    
    # 保存到文件
    output_path = "/root/.openclaw/workspace/山西代理商版本_内容.txt"
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(full_text)
    
    print(f"\n内容已保存到：{output_path}")
    
except Exception as e:
    print(f"❌ 错误：{e}")
