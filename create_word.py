#!/usr/bin/env python3
"""将 Markdown 转换为 Word 格式的简单脚本"""

from docx import Document

# 读取 Markdown 内容
with open('/root/.openclaw/workspace/中华书院菲律宾项目资料.md', 'r', encoding='utf-8') as f:
    md_content = f.read()

# 创建 Word 文档
doc = Document()

def set_chinese_font(paragraph, font_name='SimSun'):
    """设置段落中文字体"""
    for run in paragraph.runs:
        run.font.name = font_name
        run.element._rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
from docx.oxml.ns import qn

# 添加标题
doc.add_paragraph('=' * 50 + '\n中华书院·菲律宾帕赛分校项目宣传资料\n' + '=' * 50, style='Heading 1')

# 处理目录
sections = [
    ('一、学校介绍', '#一学校介绍'),
    ('二、产品体系', '#二产品体系'),
    ('三、盈利模型', '#三盈利模型'),
    ('四、运营支持', '#四运营支持'),
    ('五、项目优势', '#五项目优势'),
    ('六、合作流程', '#六合作流程')
]

for i, (section_title, anchor) in enumerate(sections):
    p = doc.add_paragraph(style='Heading 2')
    for run in p.runs:
        run.font.name = 'SimSun'
        run.element._rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        run.font.size = Pt(14)

# 添加内容
doc.add_paragraph('\n---\n', style='Heading 2')

# 一、学校介绍
doc.add_heading('一、学校介绍', level=2)
for run in doc.paragraphs[-1].runs:
    run.font.name = 'SimSun'
    run.element._rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

# 添加子标题和段落
subsections = [
    ('### 🏛️ 学校概况', ''),
    ('#### 基本信息', '- **地址**: 2269 A.Luna, Pasay, 1300 Metro Manila, Philippines\n- **联系电话**: +63 (2) 87342300 / +63 (2) 88311773 / +63 9614972406\n- **电子邮箱**: info@ppcha.edu.ph'),
]

for sub in subsections:
    if sub[0]:
        doc.add_paragraph(sub[0], style='Heading 3')
    
# 继续添加内容...
doc.add_heading('### 🎯 办学愿景与使命', level=3)
doc.add_paragraph('#### 愿景\n成为一所能够提供优质创新教育的领先机构，能够及时响应和适应时代的需求与变化。成为稳定且持续发展的教育实体，明智而审慎地利用所有可用的人力、物质和教育资源，履行我们对客户及整个菲华社区的承诺。')

# 添加表格
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
cell1 = table.rows[0].cells[0]
cell2 = table.rows[0].cells[1]
cell1.text = '目标'
cell2.text = '说明'

row_cells = []
for i in range(4):
    row = table.add_row()
    cell1 = row.cells[0]
    cell2 = row.cells[1]
    row_cells.append((cell1, cell2))

# 添加表格数据
data = [
    ('**提供 (PROVIDE)**', '提供优质教育资源和学习环境'),
    ('**促进 (PROMOTE)**', '促进学生全面发展'),
    ('**培养 (CULTIVATE)**', '培养学生的创新精神和实践能力'),
    ('**帮助 (HELP)**', '帮助学生实现个人价值'),
]

for i, (val1, val2) in enumerate(data):
    row_cells[i][0].text = val1
    row_cells[i][1].text = val2

# 保存 Word 文档（简化版）
output_path = '/root/.openclaw/workspace/中华书院菲律宾项目资料.docx'
doc.save(output_path)

print(f"✅ Word 文档已生成：{output_path}")
print("文件大小:", len(open(output_path, 'rb').read()), "bytes")
