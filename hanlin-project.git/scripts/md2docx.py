#!/usr/bin/env python3
"""Convert Markdown brochure/project description to branded DOCX."""

import sys, json, re, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def load_brand(path):
    if os.path.exists(path):
        with open(path) as f:
            return json.load(f)
    return {"colors":{"primary":"#2196F3","secondary":"#FF9800","accent":"#E3F2FD","text_dark":"#333333"},"fonts":{"heading":"Microsoft YaHei","body":"Microsoft YaHei"},"docx":{"heading1_size":22,"heading2_size":16,"body_size":12}}

def hex_to_rgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def md_to_docx(md_path, output_path, brand_path=None):
    brand = load_brand(brand_path) if brand_path else load_brand(
        os.path.join(os.path.dirname(__file__), "brand_template.json")
    )
    doc = Document()
    primary = hex_to_rgb(brand["colors"]["primary"])
    secondary = hex_to_rgb(brand["colors"]["secondary"])
    text_dark = hex_to_rgb(brand["colors"]["text_dark"])
    h1_size = brand.get("docx", {}).get("heading1_size", 22)
    h2_size = brand.get("docx", {}).get("heading2_size", 16)
    body_size = brand.get("docx", {}).get("body_size", 12)

    style = doc.styles["Normal"]
    style.font.size = Pt(body_size)
    style.font.color.rgb = RGBColor(*text_dark)

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        line = line.rstrip()
        if line.startswith("# "):
            h = doc.add_heading(line[2:], level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(*primary)
                run.font.size = Pt(h1_size)
        elif line.startswith("## "):
            h = doc.add_heading(line[3:], level=2)
            for run in h.runs:
                run.font.color.rgb = RGBColor(*primary)
                run.font.size = Pt(h2_size)
        elif line.startswith("### "):
            h = doc.add_heading(line[4:], level=3)
            for run in h.runs:
                run.font.color.rgb = RGBColor(*secondary)
        elif line.startswith("---"):
            doc.add_paragraph("").add_run("─" * 40)
        elif line.strip() == "":
            doc.add_paragraph("")
        elif line.startswith("![") and "](" in line:
            caption = line[2:line.index("]")]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"[图片: {caption}]")
            run.font.color.rgb = RGBColor(*hex_to_rgb("#888888"))
            run.font.italic = True
        else:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(body_size)

    doc.save(output_path)
    print(f"DOCX saved: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: md2docx.py <input.md> <output.docx> [brand_template.json]")
        sys.exit(1)
    md_to_docx(sys.argv[1], sys.argv[2], sys.argv[3] if len(sys.argv) > 3 else None)
